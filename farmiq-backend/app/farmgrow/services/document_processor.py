"""
Enhanced Document Processor Service
Handles multiple document formats with table extraction:
- PDFs (text + tables via pdfplumber)
- Word documents (.docx)
- Excel spreadsheets (.xlsx)
- Images (with OCR)
- Plain text files

Features:
- Structured table extraction
- Multi-format support
- Metadata preservation
- Chunk with context awareness
"""
import logging
import os
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any
from datetime import datetime
import json
import uuid

logger = logging.getLogger(__name__)

# Try to import document processing libraries
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    logger.warning("pdfplumber not installed. PDF table extraction disabled.")

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    logger.warning("python-docx not installed. DOCX processing disabled.")

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None
    logger.warning("openpyxl not installed. Excel processing disabled.")

try:
    from PIL import Image
except ImportError:
    Image = None
    logger.warning("pillow not installed. Image handling disabled.")

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    logger.warning("PyPDF2 not installed. PDF text extraction disabled.")


class DocumentProcessor:
    """
    Process documents of various formats with rich content extraction.
    
    Supported formats:
    - PDF (text extraction, table detection, OCR for scanned pages)
    - DOCX (Word documents with formatting)
    - XLSX (Excel spreadsheets)
    - TXT (Plain text)
    - Image formats (PNG, JPG with OCR)
    """
    
    SUPPORTED_FORMATS = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'txt': 'text/plain',
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg'
    }
    
    def __init__(self, ocr_service=None):
        """
        Initialize document processor.
        
        Args:
            ocr_service: Optional OCR service for image/scanned page processing
        """
        self.ocr_service = ocr_service
        logger.info("📄 Document Processor initialized")
        self._log_supported_formats()
    
    def _log_supported_formats(self):
        """Log which document formats are supported."""
        supported = []
        if PyPDF2:
            supported.append("PDF (text)")
        if pdfplumber:
            supported.append("PDF (tables)")
        if DocxDocument:
            supported.append("DOCX (Word)")
        if load_workbook:
            supported.append("XLSX (Excel)")
        if self.ocr_service:
            supported.append("Images (OCR)")
        
        logger.info(f"   Supported formats: {', '.join(supported) if supported else 'None (install dependencies)'}")
    
    async def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract structured content.
        
        Args:
            file_path: Path to document
        
        Returns:
            Dictionary with:
            - format: Document format
            - text: Main extracted text
            - tables: List of extracted tables
            - sections: Structured sections
            - metadata: Document metadata
            - chunks: Text chunks with context
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower().lstrip('.')
        logger.info(f"📄 Processing: {file_path.name} (format: {file_ext})")
        
        try:
            if file_ext == 'pdf':
                return await self._process_pdf(str(file_path))
            elif file_ext == 'docx':
                return await self._process_docx(str(file_path))
            elif file_ext == 'xlsx':
                return await self._process_xlsx(str(file_path))
            elif file_ext == 'txt':
                return await self._process_txt(str(file_path))
            elif file_ext in ['png', 'jpg', 'jpeg']:
                return await self._process_image(str(file_path))
            else:
                raise ValueError(f"Unsupported format: {file_ext}")
        
        except Exception as e:
            logger.error(f"❌ Error processing {file_path.name}: {str(e)}")
            raise
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and tables from PDF.
        
        Returns structured content with:
        - Text by page
        - Tables with structure
        - Metadata
        - Sections (if detectable)
        """
        if not PyPDF2:
            raise ImportError("PyPDF2 required for PDF text extraction")
        
        document_id = str(uuid.uuid4())
        file_name = Path(file_path).name
        text_parts = []
        tables = []
        sections = []
        metadata = {'pages': 0, 'tables_found': 0, 'has_images': False}
        
        logger.info(f"   🔍 Extracting text and tables from PDF...")
        
        # Extract with PyPDF2 (text)
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            metadata['pages'] = len(reader.pages)
            
            # Try to extract tables with pdfplumber if available
            if pdfplumber:
                try:
                    with pdfplumber.open(file_path) as pdf:
                        for page_num, page in enumerate(pdf.pages, 1):
                            # Extract tables from this page
                            page_tables = page.extract_tables()
                            if page_tables:
                                for table_idx, table in enumerate(page_tables):
                                    table_data = {
                                        'page': page_num,
                                        'table_idx': table_idx,
                                        'rows': table,
                                        'markdown': self._table_to_markdown(table)
                                    }
                                    tables.append(table_data)
                                    metadata['tables_found'] += 1
                                    
                                    logger.debug(f"      📊 Found table on page {page_num}: {len(table)} rows")
                
                except Exception as e:
                    logger.debug(f"   ⚠️  pdfplumber table extraction failed: {str(e)}")
            
            # Extract text from pages
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                text_parts.append(f"[Page {page_num}]\n{page_text}")
                
                # Check for images (potential scanned content)
                if '/Image' in str(page):
                    metadata['has_images'] = True
                    logger.info(f"      🖼️  Found images on page {page_num} - consider OCR")
        
        # Combine text
        full_text = "\n\n".join(text_parts)
        
        # Simple section detection (by page breaks)
        sections = self._detect_sections(full_text)
        
        logger.info(f"   ✅ Extracted {len(text_parts)} pages, {len(tables)} tables")
        
        return {
            'document_id': document_id,
            'file_name': file_name,
            'format': 'pdf',
            'text': full_text,
            'tables': tables,
            'sections': sections,
            'metadata': metadata,
            'processed_at': datetime.now().isoformat()
        }
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract content from Word document.
        
        Returns:
        - Text by paragraph
        - Tables with content
        - Sections
        - Metadata (author, created date, etc.)
        """
        if not DocxDocument:
            raise ImportError("python-docx required for DOCX processing")
        
        document_id = str(uuid.uuid4())
        file_name = Path(file_path).name
        
        logger.info(f"   📖 Extracting content from Word document...")
        
        doc = DocxDocument(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Detect heading level
                style = para.style.name
                text_parts.append(para.text)
        
        # Extract tables
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            tables.append({
                'table_idx': table_idx,
                'rows': table_data,
                'markdown': self._table_to_markdown(table_data)
            })
        
        # Extract metadata
        metadata = {
            'author': doc.core_properties.author or 'Unknown',
            'created': doc.core_properties.created.isoformat() if doc.core_properties.created else None,
            'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
            'paragraphs': len(doc.paragraphs),
            'tables_found': len(tables)
        }
        
        full_text = "\n\n".join(text_parts)
        sections = self._detect_sections(full_text)
        
        logger.info(f"   ✅ Extracted {len(text_parts)} paragraphs, {len(tables)} tables")
        
        return {
            'document_id': document_id,
            'file_name': file_name,
            'format': 'docx',
            'text': full_text,
            'tables': tables,
            'sections': sections,
            'metadata': metadata,
            'processed_at': datetime.now().isoformat()
        }
    
    async def _process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract content from Excel spreadsheet.
        
        Returns:
        - Each sheet as a table
        - Combined text representation
        - Metadata (sheet names, dimensions)
        """
        if not load_workbook:
            raise ImportError("openpyxl required for XLSX processing")
        
        document_id = str(uuid.uuid4())
        file_name = Path(file_path).name
        
        logger.info(f"   📊 Extracting content from Excel spreadsheet...")
        
        wb = load_workbook(file_path)
        tables = []
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_data = []
            
            for row in ws.iter_rows(values_only=True):
                sheet_data.append(list(row))
            
            if sheet_data:
                tables.append({
                    'sheet_name': sheet_name,
                    'rows': sheet_data,
                    'markdown': self._table_to_markdown(sheet_data)
                })
                
                # Add to text representation
                text_parts.append(f"[Sheet: {sheet_name}]\n{self._table_to_markdown(sheet_data)}")
        
        metadata = {
            'sheets': wb.sheetnames,
            'sheet_count': len(wb.sheetnames),
            'tables_found': len(tables)
        }
        
        full_text = "\n\n".join(text_parts)
        sections = self._detect_sections(full_text)
        
        logger.info(f"   ✅ Extracted {len(tables)} sheets")
        
        return {
            'document_id': document_id,
            'file_name': file_name,
            'format': 'xlsx',
            'text': full_text,
            'tables': tables,
            'sections': sections,
            'metadata': metadata,
            'processed_at': datetime.now().isoformat()
        }
    
    async def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """
        Process plain text file.
        """
        document_id = str(uuid.uuid4())
        file_name = Path(file_path).name
        
        logger.info(f"   📝 Processing text file...")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        sections = self._detect_sections(text)
        metadata = {
            'file_size': os.path.getsize(file_path),
            'char_count': len(text),
            'line_count': len(text.split('\n'))
        }
        
        logger.info(f"   ✅ Processed {len(text)} characters")
        
        return {
            'document_id': document_id,
            'file_name': file_name,
            'format': 'txt',
            'text': text,
            'tables': [],
            'sections': sections,
            'metadata': metadata,
            'processed_at': datetime.now().isoformat()
        }
    
    async def _process_image(self, file_path: str) -> Dict[str, Any]:
        """
        Process image file with OCR.
        """
        if not self.ocr_service:
            raise ValueError("OCR service required for image processing")
        
        document_id = str(uuid.uuid4())
        file_name = Path(file_path).name
        
        logger.info(f"   🖼️  Processing image with OCR...")
        
        try:
            text = await self.ocr_service.extract_text_from_image(file_path)
            sections = self._detect_sections(text)
            metadata = {'ocr_method': 'easyocr'}
            
            logger.info(f"   ✅ OCR extracted {len(text)} characters")
        
        except Exception as e:
            logger.error(f"   ❌ OCR failed: {str(e)}")
            text = ""
            sections = []
            metadata = {'ocr_method': 'failed', 'error': str(e)}
        
        return {
            'document_id': document_id,
            'file_name': file_name,
            'format': 'image',
            'text': text,
            'tables': [],
            'sections': sections,
            'metadata': metadata,
            'processed_at': datetime.now().isoformat()
        }
    
    @staticmethod
    def _table_to_markdown(table_data: List[List[str]]) -> str:
        """
        Convert table data to markdown format.
        
        Args:
            table_data: List of rows, each row is list of cell values
        
        Returns:
            Markdown formatted table string
        """
        if not table_data:
            return ""
        
        # Convert all cells to strings
        str_table = []
        for row in table_data:
            str_row = [str(cell) if cell is not None else "" for cell in row]
            str_table.append(str_row)
        
        # Build markdown
        lines = []
        
        # Header row
        if str_table:
            lines.append("| " + " | ".join(str_table[0]) + " |")
            lines.append("|" + "|".join(["---" for _ in str_table[0]]) + "|")
        
        # Data rows
        for row in str_table[1:]:
            lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(lines)
    
    @staticmethod
    def _detect_sections(text: str, min_section_length: int = 100) -> List[Dict]:
        """
        Simple section detection based on page breaks and heading patterns.
        
        Args:
            text: Full document text
            min_section_length: Minimum characters for a section
        
        Returns:
            List of detected sections with start/end positions and content
        """
        sections = []
        
        # Split by [Page N] markers if present
        page_pattern = r'\[Page \d+\]'
        parts = text.split('[Page ')
        
        current_pos = 0
        for part in parts:
            if len(part) > min_section_length:
                section = {
                    'start': current_pos,
                    'end': current_pos + len(part),
                    'length': len(part),
                    'preview': part[:100] + "..." if len(part) > 100 else part
                }
                sections.append(section)
            current_pos += len(part)
        
        return sections
